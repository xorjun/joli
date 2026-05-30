import os
import subprocess
import uuid
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "outputs")


def _ensure_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(os.path.join(OUTPUT_DIR, "photos"), exist_ok=True)
    os.makedirs(os.path.join(OUTPUT_DIR, "zeugnisse"), exist_ok=True)


def _set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr, val in edge_data.items():
                element.set(qn(f'w:{attr}'), str(val))
            tc_borders.append(element)
    tc_pr.append(tc_borders)


def _add_section_line(doc, width_pts: float = 0.5):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(width_pts * 8)))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def generate_resume_docx(profile: dict, language: str = "de") -> str:
    """Generate a DIN 5008 tabellarischer Lebenslauf as DOCX."""
    _ensure_output_dir()
    doc = Document()

    # Page setup: A4, DIN 5008 Form B margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.67)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)

    is_german = language == "de"

    # ---- Header: Name + Personal Data ----
    name = profile.get("full_name", "") or "Vorname Nachname"
    p = doc.add_paragraph()
    run = p.add_run(name)
    _set_font(run, size=18, bold=True, color=(0x33, 0x33, 0x33))

    # Personal data line
    personal_parts = []
    if profile.get("street"):
        personal_parts.append(profile["street"])
    if profile.get("postal_code") or profile.get("city"):
        personal_parts.append(f"{profile.get('postal_code', '')} {profile.get('city', '')}".strip())
    if profile.get("phone"):
        personal_parts.append(f"Tel: {profile['phone']}")
    if profile.get("email"):
        personal_parts.append(profile["email"])

    if personal_parts:
        p = doc.add_paragraph()
        run = p.add_run(" · ".join(personal_parts))
        _set_font(run, size=9, color=(0x66, 0x66, 0x66))

    if is_german:
        p = doc.add_paragraph()
        pers_data = []
        if profile.get("date_of_birth"):
            pers_data.append(f"Geboren: {profile['date_of_birth']}")
        if profile.get("place_of_birth"):
            pers_data.append(f"in {profile['place_of_birth']}")
        if profile.get("nationality"):
            pers_data.append(f"Staatsangehörigkeit: {profile['nationality']}")
        if profile.get("marital_status"):
            pers_data.append(f"Familienstand: {profile['marital_status']}")
        if pers_data:
            run = p.add_run(" · ".join(pers_data))
            _set_font(run, size=9, color=(0x66, 0x66, 0x66))

    # ---- Berufserfahrung ----
    _add_section_line(doc)
    h = doc.add_paragraph()
    run = h.add_run("Berufserfahrung" if is_german else "Work Experience")
    _set_font(run, size=14, bold=True, color=(0x33, 0x33, 0x33))

    experiences = profile.get("work_experiences", [])
    for exp in experiences:
        date_str = _format_date_range(exp.get("start_date", ""), exp.get("end_date", ""), exp.get("is_current", False), is_german)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Date cell
        cell_date = table.cell(0, 0)
        cell_date.width = Cm(3.0)
        p = cell_date.paragraphs[0]
        run = p.add_run(date_str)
        _set_font(run, size=10, color=(0x66, 0x66, 0x66))

        # Content cell
        cell_content = table.cell(0, 1)
        p = cell_content.paragraphs[0]
        run = p.add_run(exp.get("title", ""))
        _set_font(run, size=11, bold=True)
        if exp.get("company"):
            p = cell_content.add_paragraph()
            company_loc = exp.get("company", "")
            if exp.get("company_location"):
                company_loc += f", {exp['company_location']}"
            run = p.add_run(company_loc)
            _set_font(run, size=10, color=(0x66, 0x66, 0x66))

        # Achievements
        achievements = exp.get("achievements", [])
        if isinstance(achievements, str):
            achievements = [achievements]
        for ach in achievements[:5]:
            p = cell_content.add_paragraph()
            p.style = doc.styles['List Bullet']
            run = p.add_run(str(ach))
            _set_font(run, size=10)

        # Description if no achievements
        if not achievements and exp.get("description_md"):
            p = cell_content.add_paragraph()
            run = p.add_run(exp["description_md"][:300])
            _set_font(run, size=10)

    # ---- Ausbildung ----
    educations = profile.get("educations", [])
    if educations:
        _add_section_line(doc)
        h = doc.add_paragraph()
        run = h.add_run("Ausbildung" if is_german else "Education")
        _set_font(run, size=14, bold=True, color=(0x33, 0x33, 0x33))

        for edu in educations:
            sy = edu.get("start_year", "")
            ey = edu.get("end_year", "")
            date_str = f"{sy} – {ey}" if sy and ey else str(sy or ey or "")

            table = doc.add_table(rows=1, cols=2)
            cell_date = table.cell(0, 0)
            cell_date.width = Cm(3.0)
            p = cell_date.paragraphs[0]
            run = p.add_run(date_str)
            _set_font(run, size=10, color=(0x66, 0x66, 0x66))

            cell_content = table.cell(0, 1)
            p = cell_content.paragraphs[0]
            degree_text = edu.get("degree", "")
            if edu.get("field"):
                degree_text += f" {edu['field']}"
            run = p.add_run(degree_text)
            _set_font(run, size=11, bold=True)

            inst_text = edu.get("institution", "")
            if edu.get("institution_location"):
                inst_text += f", {edu['institution_location']}"
            if inst_text:
                p = cell_content.add_paragraph()
                run = p.add_run(inst_text)
                _set_font(run, size=10, color=(0x66, 0x66, 0x66))

            if edu.get("grade"):
                p = cell_content.add_paragraph()
                run = p.add_run(f"Note: {edu['grade']}" if is_german else f"Grade: {edu['grade']}")
                _set_font(run, size=10, italic=True)

    # ---- Kenntnisse & Fähigkeiten ----
    _add_section_line(doc)
    h = doc.add_paragraph()
    run = h.add_run("Kenntnisse & Fähigkeiten" if is_german else "Skills")
    _set_font(run, size=14, bold=True, color=(0x33, 0x33, 0x33))

    # Language skills
    lang_skills = profile.get("language_skills", [])
    if lang_skills:
        p = doc.add_paragraph()
        run = p.add_run("Sprachkenntnisse:" if is_german else "Languages:")
        _set_font(run, size=10, bold=True)
        lang_texts = []
        for ls in lang_skills:
            level = ls.get("cefr_level", "")
            lang_texts.append(f"{ls['language']} ({level})")
        p = doc.add_paragraph()
        run = p.add_run(", ".join(lang_texts))
        _set_font(run, size=10)

    # Tech skills
    tech_skills = profile.get("tech_skills", [])
    if tech_skills:
        p = doc.add_paragraph()
        run = p.add_run("IT-Kenntnisse:" if is_german else "Technical:")
        _set_font(run, size=10, bold=True)
        p = doc.add_paragraph()
        skill_names = [s["name"] for s in tech_skills]
        run = p.add_run(", ".join(skill_names))
        _set_font(run, size=10)

    # ---- Footer ----
    _add_section_line(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    city = profile.get("city", "")
    date_str = "30.05.2026"  # Will be set dynamically in production
    run = p.add_run(f"{city}, {date_str}" if city else date_str)
    _set_font(run, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("_____________")
    _set_font(run, size=10, color=(0x99, 0x99, 0x99))

    # Save
    filename = f"{uuid.uuid4()}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


def generate_cover_letter_docx(profile: dict, job: dict, body_md: str, language: str = "de") -> str:
    """Generate a DIN 5008 Anschreiben as DOCX."""
    _ensure_output_dir()
    doc = Document()

    # Page setup: A4, DIN 5008 Form B
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.67)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)

    is_german = language == "de"

    # Sender block (compact, single line style)
    name = profile.get("full_name", "")
    sender_parts = [name] if name else []
    if profile.get("street"):
        sender_parts.append(profile["street"])
    if profile.get("postal_code") or profile.get("city"):
        sender_parts.append(f"{profile.get('postal_code', '')} {profile.get('city', '')}".strip())
    if profile.get("phone"):
        sender_parts.append(f"Tel: {profile['phone']}")
    if profile.get("email"):
        sender_parts.append(profile["email"])

    p = doc.add_paragraph()
    run = p.add_run(" · ".join(sender_parts))
    _set_font(run, size=8, color=(0x33, 0x33, 0x33))

    # Spacer for window envelope
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)

    # Recipient block
    company = job.get("company", "")
    if company:
        p = doc.add_paragraph()
        run = p.add_run(company)
        _set_font(run, size=10, bold=True)

    p = doc.add_paragraph()
    recipient_text = f"{job.get('company_location', '')}"
    run = p.add_run(recipient_text)
    _set_font(run, size=10)

    # Date
    city = profile.get("city", "")
    date_str = "30.05.2026"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{city}, {date_str}" if city else date_str)
    _set_font(run, size=10)

    # Subject line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    subject = f"Bewerbung als {job.get('job_title', '')}" if is_german else f"Application for {job.get('job_title', '')}"
    if job.get("reference_number"):
        subject += f" — Kennziffer: {job['reference_number']}"
    run = p.add_run(subject)
    _set_font(run, size=11, bold=True)

    # Salutation
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Sehr geehrte Damen und Herren," if is_german else "Dear Hiring Manager,")
    _set_font(run, size=11)

    # Body
    for line in body_md.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("#"):
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        _set_font(run, size=11)

    # Closing
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run("Mit freundlichen Grüßen" if is_german else "Sincerely,")
    _set_font(run, size=11)

    # Signature placeholder
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run(name or "_____________")
    _set_font(run, size=11)

    # Anlagen
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run("Anlagen:" if is_german else "Enclosures:")
    _set_font(run, size=8, bold=True)

    anlagen = ["Lebenslauf" if is_german else "Resume"]
    zeugnisse = profile.get("zeugnisse", [])
    for z in zeugnisse[:5]:
        anlagen.append(z.get("title", "Zeugnis"))
    certificates = profile.get("certificates", [])
    for c in certificates[:5]:
        anlagen.append(c.get("name", "Zertifikat"))

    for a in anlagen:
        p = doc.add_paragraph()
        run = p.add_run(f"  – {a}")
        _set_font(run, size=8, color=(0x66, 0x66, 0x66))

    # Save
    filename = f"{uuid.uuid4()}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


def convert_to_pdf(docx_path: str) -> str | None:
    """Convert DOCX to PDF using LibreOffice headless."""
    try:
        output_dir = os.path.dirname(docx_path)
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", output_dir],
            capture_output=True, text=True, timeout=30,
        )
        base = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(output_dir, f"{base}.pdf")
        if os.path.exists(pdf_path):
            return pdf_path
    except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
        pass
    return None


def _format_date_range(start: str, end: str, is_current: bool, is_german: bool) -> str:
    """Format MM/YYYY to MM.YYYY display."""
    def fmt(d):
        if not d:
            return ""
        parts = d.replace("-", "/").replace(".", "/").split("/")
        if len(parts) == 2:
            return f"{parts[0]}/{parts[1]}"
        return d

    start_fmt = fmt(start)
    if is_current:
        end_fmt = "heute" if is_german else "present"
    else:
        end_fmt = fmt(end)
    return f"{start_fmt} – {end_fmt}"
